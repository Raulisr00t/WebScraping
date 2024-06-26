import data_NVD


from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import exploit_code

# This function creates our report in word format
def create_docx(cve_id):
    summary, cvss_score, vector, reference ,  reference1, reference2, reference3, reference4, cve_link, published_date, last_modified_date, source = data_NVD.get_detail(cve_id)
    exploit_number, exploit_header, exploit_script = exploit_code.check_exploit_db(cve_id)

    file = Document()

    # Title
    title = file.add_heading('CVE Information', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CVE details in word format 
    cve_details = file.add_paragraph()
    cve_details.add_run("CVE ID: ").bold = True
    cve_details.add_run(f"{cve_id}\n\n")

    cve_details.add_run("Summary: ").bold = True
    cve_details.add_run(f"{summary}\n\n")

    cve_details.add_run("Base Score: ").bold = True
    cve_details.add_run(f"{cvss_score}\n\n")

    cve_details.add_run("Vector: ").bold = True
    cve_details.add_run(f"{vector}\n\n")

    cve_details.add_run("Reference: ").bold = True
    cve_details.add_run(f"{reference}\n\n")
    cve_details.add_run(f"{reference1}\n\n")
    cve_details.add_run(f"{reference2}\n\n")
    cve_details.add_run(f"{reference3}\n\n")
    cve_details.add_run(f"{reference4}\n\n")
    

    cve_details.add_run("Change History: ").bold = True
    cve_details.add_run(f"{cve_link}\n\n")
    cve_details.add_run(f"{published_date}\n\n")
    cve_details.add_run(f"{last_modified_date}\n\n")
    cve_details.add_run(f"{source}\n\n")  
    

    # Add spacing
    cve_details.add_run("\n" * 2)

    # Exploit Information
    exploit_paragraph = file.add_paragraph()
    exploit_paragraph.add_run("Number of exploit scripts were found: ").bold = True
    exploit_paragraph.add_run(f"{exploit_number}\n\n")
    exploit_paragraph.add_run(f"\nExploit Header: \n\n")
    exploit_paragraph.add_run(f"{exploit_header}\n\n")

    exploit_script_lines = exploit_script.split('\n')
    for line in exploit_script_lines:
        exploit_paragraph.add_run(line).font.size = Pt(10)
        exploit_paragraph.add_run('\n')



    #save file as doc format
    docx_filename = f"report_{cve_id}.docx"
    file.save(docx_filename)

    return docx_filename

# This function creates report in pdf 
def create_pdf(cve_id):
    summary, cvss_score, vector, reference ,  reference1, reference2, reference3, reference4 , cve_link, published_date, last_modified_data, source = data_NVD.get_detail(cve_id)
    exploit_number, exploit_header, exploit_script = exploit_code.check_exploit_db(cve_id)

    pdf_file_name = f'report_{cve_id}.pdf'
    doc = SimpleDocTemplate(pdf_file_name, pagesize=letter)

    data = []

    #form of style type
    styles = getSampleStyleSheet()
    style_heading = styles['Heading1']
    style_normal = styles['Normal']
    
    # part of create data about the CVE
    data.append(Paragraph('CVE Information', style_heading))
    data.append(Paragraph(f'<b>CVE ID:</b> {cve_id}\n', style_normal))
    data.append(Paragraph(f'<b>Summary:</b> {summary}\n', style_normal))
    data.append(Paragraph(f'<b>Base Score:</b> <p color="red" style="height:50px">{cvss_score}</p>\n', style_normal))
    data.append(Paragraph(f'<b>Vector:</b> {vector}\n', style_normal))
    data.append(Paragraph(f'<b>Reference:</b> <a color="blue" href="{reference}">{reference}</a>\n', style_normal))
    data.append(Paragraph(f'<b>Reference1:</b> <a color="blue" href="{reference1}">{reference1}</a>\n', style_normal))
    data.append(Paragraph(f'<b>Reference2:</b> <a color="blue" href="{reference2}">{reference2}</a>\n', style_normal))
    data.append(Paragraph(f'<b>Reference3:</b> <a color="blue" href="{reference3}">{reference3}</a>\n', style_normal))
    data.append(Paragraph(f'<b>Reference4:</b> <a color="blue" href="{reference4}">{reference4}</a>\n', style_normal))
    data.append(Paragraph('<br/><br/>', style_normal))
    data.append(Paragraph(f'<b>Cve_Link:</b> {cve_link}\n', style_normal))
    data.append(Paragraph(f'<b>Last Modified data:</b> {last_modified_data}\n', style_normal))
    data.append(Paragraph(f'<b>Source:</b> {source}\n', style_normal))
    data.append(Paragraph(f'<b>Published Data:</b> {published_date}\n', style_normal))
    
    #part of exploit code in pdf
    data.append(Paragraph('Exploit Information', style_heading))
    data.append(Paragraph(f'<b>Number of exploit scripts found:</b> {exploit_number}', style_normal))
    data.append(Paragraph(f'<b>Exploit Header:</b> {exploit_header}', style_normal))
    data.append(Paragraph(f'<b>Exploit Script:</b>', style_normal))

    exploit_script_lines = exploit_script.split('\n')
    for line in exploit_script_lines:
        data.append(Paragraph(f'<font face="Courier">{line}</font>', style_normal))

    doc.build(data)

    return pdf_file_name



# This part of report creates report in md format
def create_md(cve_id):
    summary, cvss_score, vector, reference,  reference1, reference2, reference3, reference4 , cve_link, published_date, last_modified_data, source  = data_NVD.get_detail(cve_id)
    exploit_number, exploit_header, exploit_script = exploit_code.check_exploit_db(cve_id)

    md_content = ['## CVE Information', f'**CVE ID:** {cve_id}', f'**Summary:** {summary}',
                  f'**Base Score:** {cvss_score}', f'**Vector:** {vector}',
                  f'**Reference:** [{reference}]({reference})\n',
                  f'**Reference1:** [{reference1}]({reference1})\n',
                  f'**Reference2:** [{reference2}]({reference2})\n',
                  f'**Reference3:** [{reference3}]({reference3})\n',
                  f'**Reference4:** [{reference4}]({reference4})\n',
                  f'**CVE_Link:** [{cve_link}]({cve_link})\n',
                  f'**Last_Modified_data:** [{last_modified_data}]({last_modified_data})\n',
                  f'**Published_data:** [{published_date}]({published_date})\n',
                  f'**Source:** [{source}]({source})\n',
                  f'**Number of exploit scripts found:** {exploit_number}', f'**Exploit Header:** {exploit_header}',
                  '**Exploit Script:**'
                    ]
    exploit_script_lines = exploit_script.split('\n')
    md_content.extend(['```', *exploit_script_lines, '```'])
  

    md_filename = f"report_{cve_id}.md"
    with open(md_filename, 'w', encoding='utf-8') as md_file:
        md_file.write('\n'.join(md_content))

    return md_filename
